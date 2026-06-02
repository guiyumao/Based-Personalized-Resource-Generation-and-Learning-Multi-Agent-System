# -*- coding: utf-8 -*-
"""Generate the final Word document with database tables, ER diagram, and flowchart."""

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_DOCX = DOCS_DIR / "agent-main-merge-summary.docx"
ER_IMAGE = DOCS_DIR / "evaluation-er-diagram.png"
FLOW_IMAGE = DOCS_DIR / "evaluation-agent-flowchart.png"

BG = "#F6F3EE"
INK = "#1F2937"
MUTED = "#475569"
LINE = "#8B5E3C"
ACCENT = "#C96A3D"
ENTITY_FILL = "#FFFDF8"
ENTITY_HEADER = "#EADBC8"
ENTITY_BORDER = "#B08968"
LABEL_FILL = "#FFFFFF"
LEGEND_FILL = "#FDF3E7"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """Load a Chinese-capable system font."""

    candidates = (
        [
            r"C:\Windows\Fonts\msyhbd.ttc",
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
        ]
        if bold
        else [
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\msyhbd.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
        ]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


FONT_TITLE = load_font(42, bold=True)
FONT_SUB = load_font(24)
FONT_HEADER = load_font(24, bold=True)
FONT_BODY = load_font(19)
FONT_SMALL = load_font(18)
FONT_LABEL = load_font(18, bold=True)
FONT_MONO = load_font(20)


def rounded_box(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, *, fill: str, outline: str) -> None:
    """Draw a rounded box."""

    draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=fill, outline=outline, width=3)


def draw_entity(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, title: str, fields: list[str]) -> None:
    """Draw one entity box."""

    rounded_box(draw, x, y, w, h, fill=ENTITY_FILL, outline=ENTITY_BORDER)
    draw.rounded_rectangle((x, y, x + w, y + 44), radius=22, fill=ENTITY_HEADER, outline=ENTITY_BORDER, width=3)
    draw.rectangle((x, y + 22, x + w, y + 44), fill=ENTITY_HEADER, outline=ENTITY_HEADER)
    draw.text((x + 16, y + 9), title, fill=INK, font=FONT_HEADER)
    current_y = y + 60
    for field in fields:
        draw.text((x + 18, current_y), field, fill=INK, font=FONT_BODY)
        current_y += 28


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    *,
    start_label: str | None = None,
    end_label: str | None = None,
) -> None:
    """Draw a connector with optional cardinality labels."""

    draw.line(points, fill=LINE, width=4)
    x2, y2 = points[-1]
    x1, y1 = points[-2]
    dx, dy = x2 - x1, y2 - y1
    if abs(dx) > abs(dy):
        sign = 1 if dx > 0 else -1
        arrow = [(x2, y2), (x2 - 16 * sign, y2 - 8), (x2 - 16 * sign, y2 + 8)]
    else:
        sign = 1 if dy > 0 else -1
        arrow = [(x2, y2), (x2 - 8, y2 - 16 * sign), (x2 + 8, y2 - 16 * sign)]
    draw.polygon(arrow, fill=LINE)
    if start_label:
        draw_cardinality(draw, points[0], start_label)
    if end_label:
        draw_cardinality(draw, points[-1], end_label)


def draw_cardinality(draw: ImageDraw.ImageDraw, point: tuple[int, int], text: str) -> None:
    """Draw one relationship cardinality label."""

    x, y = point
    bbox = draw.textbbox((0, 0), text, font=FONT_LABEL)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    pad = 8
    draw.rounded_rectangle(
        (x - text_w // 2 - pad, y - text_h // 2 - pad, x + text_w // 2 + pad, y + text_h // 2 + pad),
        radius=10,
        fill=LABEL_FILL,
        outline=ACCENT,
        width=2,
    )
    draw.text((x - text_w // 2, y - text_h // 2 - 1), text, fill=ACCENT, font=FONT_LABEL)


def entity_anchor(spec: dict[str, object], position: str) -> tuple[int, int]:
    """Return anchor point for an entity."""

    x, y = spec["xy"]
    w, h = spec["size"]
    x = int(x)
    y = int(y)
    w = int(w)
    h = int(h)
    if position == "left":
        return (x, y + h // 2)
    if position == "right":
        return (x + w, y + h // 2)
    if position == "top":
        return (x + w // 2, y)
    return (x + w // 2, y + h)


def generate_er_diagram() -> None:
    """Generate the formal ER diagram image."""

    width, height = 2400, 1700
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    blue = "#4C78D0"
    blue_dark = "#2F5FB3"
    blue_light = "#EAF1FF"
    warn = "#FFE9B5"
    warn_border = "#E6BF58"

    def draw_note_box(x: int, y: int, w: int, h: int, fill: str, border: str, text: str) -> None:
        draw.rectangle((x, y, x + w, y + h), fill=fill, outline=border, width=2)
        draw.text((x + 18, y + 10), text, fill=INK, font=FONT_SMALL)

    def draw_blue_entity(x: int, y: int, w: int, h: int, title: str, subtitle: str | None = None) -> None:
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=blue, outline=blue_dark, width=3)
        bbox = draw.textbbox((0, 0), title, font=FONT_HEADER)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.text((x + (w - tw) / 2, y + (h - th) / 2 - 8), title, fill="white", font=FONT_HEADER)
        if subtitle:
            sb = draw.textbbox((0, 0), subtitle, font=FONT_SMALL)
            stw = sb[2] - sb[0]
            draw.text((x + (w - stw) / 2, y + h - 34), subtitle, fill="#EAF1FF", font=FONT_SMALL)

    def draw_relation_diamond(cx: int, cy: int, w: int, h: int, text: str) -> None:
        points = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
        draw.polygon(points, fill=blue_light, outline=blue_dark)
        bbox = draw.textbbox((0, 0), text, font=FONT_SMALL)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.text((cx - tw / 2, cy - th / 2 - 1), text, fill=blue_dark, font=FONT_SMALL)

    def draw_line_with_labels(x1: int, y1: int, x2: int, y2: int, start_label: str | None = None, end_label: str | None = None) -> None:
        draw.line((x1, y1, x2, y2), fill=blue, width=3)
        if start_label:
            draw_cardinality(draw, (x1, y1), start_label)
        if end_label:
            draw_cardinality(draw, (x2, y2), end_label)

    # heading like the reference document
    draw.text((110, 55), "1.4  E-R 图", fill=blue_dark, font=FONT_TITLE)
    draw_note_box(
        110,
        125,
        2080,
        44,
        blue_light,
        blue,
        "使用课程设计风格展示系统核心数据实体，并标注主要的 1:1、1:N、N:N 关系。",
    )
    draw_note_box(
        110,
        178,
        2080,
        44,
        warn,
        warn_border,
        "E-R 图应覆盖评估智能体所依赖的真实答题、用户画像、题目、知识点、报告和学习路径等核心数据对象。",
    )

    # entities
    draw_blue_entity(150, 350, 240, 92, "用户", "users")
    draw_blue_entity(520, 250, 280, 92, "用户画像", "user_profiles")
    draw_blue_entity(930, 250, 260, 92, "知识点", "knowledge_points")
    draw_blue_entity(1320, 250, 260, 92, "题目", "exercises")
    draw_blue_entity(1720, 350, 300, 92, "答题记录", "answer_records")
    draw_blue_entity(930, 540, 300, 92, "学习报告", "learning_reports")
    draw_blue_entity(520, 780, 280, 92, "学习路径", "learning_paths")
    draw_blue_entity(930, 780, 280, 92, "学习任务", "learning_tasks")
    draw_blue_entity(1320, 780, 260, 92, "学习资源", "resources")
    draw_blue_entity(1720, 780, 320, 92, "知识点关系", "knowledge_relations")

    # relations
    draw_relation_diamond(455, 315, 120, 70, "拥有")
    draw_relation_diamond(865, 315, 120, 70, "关联")
    draw_relation_diamond(1650, 315, 120, 70, "作答")
    draw_relation_diamond(1450, 560, 120, 70, "生成")
    draw_relation_diamond(865, 825, 120, 70, "包含")
    draw_relation_diamond(1265, 825, 120, 70, "指向")
    draw_relation_diamond(1655, 825, 140, 70, "关联 / 依赖")

    # connectors and cardinalities
    draw_line_with_labels(390, 396, 455, 340, "1", "1")
    draw_line_with_labels(520, 296, 455, 340, "1", "1")

    draw_line_with_labels(800, 296, 865, 340, "1", "N")
    draw_line_with_labels(930, 296, 865, 340, "1", "N")

    draw_line_with_labels(1190, 296, 1260, 296, "1", "N")
    draw_line_with_labels(1260, 296, 1320, 296, "1", "N")

    draw_line_with_labels(1580, 296, 1650, 340, "1", "N")
    draw_line_with_labels(1720, 396, 1650, 340, "N", "1")

    draw_line_with_labels(1870, 442, 1870, 710, "N", "N")
    draw_line_with_labels(1180, 586, 1390, 560, "1", "N")
    draw_line_with_labels(1510, 560, 1860, 560, "1", "N")

    draw_line_with_labels(800, 826, 865, 826, "1", "N")
    draw_line_with_labels(930, 826, 865, 826, "1", "N")

    draw_line_with_labels(1210, 826, 1265, 826, "1", "N")
    draw_line_with_labels(1320, 826, 1265, 826, "1", "N")

    draw_line_with_labels(1580, 826, 1655, 826, "1", "N")
    draw_line_with_labels(1720, 826, 1655, 826, "N", "N")

    # report and path links
    draw_line_with_labels(660, 442, 660, 735, "1", "N")
    draw_line_with_labels(1505, 442, 1505, 735, "1", "N")

    # legend
    draw.rectangle((140, 1080, 2140, 1500), fill="#FAFBFF", outline=blue, width=2)
    draw.text((170, 1110), "图示说明", fill=blue_dark, font=FONT_HEADER)
    legend_lines = [
        "1. 蓝色圆角矩形表示核心数据实体，中文名称用于课程设计展示，英文名称用于与数据库表对应。",
        "2. 菱形表示实体之间的业务联系，如“拥有”“作答”“生成”“包含”等。",
        "3. 连线两端的 1 / N 表示基数关系，体现一对一、一对多和多对多结构。",
        "4. 该图重点围绕评估智能体主链路：用户 -> 题目 -> 答题记录 -> 用户画像 / 学习报告。",
        "5. learning_paths、learning_tasks、resources、knowledge_relations 为评估结果向路径调整和资源推荐扩展提供支撑。",
    ]
    current_y = 1160
    for line_text in legend_lines:
        draw.text((175, current_y), line_text, fill=MUTED, font=FONT_SMALL)
        current_y += 48

    image.save(ER_IMAGE)


def draw_flow_box(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, title: str, lines: list[str]) -> None:
    """Draw one process box in the flowchart."""

    rounded_box(draw, x, y, w, h, fill=ENTITY_FILL, outline=ENTITY_BORDER)
    draw.rounded_rectangle((x, y, x + w, y + 46), radius=24, fill=ENTITY_HEADER, outline=ENTITY_BORDER, width=4)
    draw.rectangle((x, y + 24, x + w, y + 46), fill=ENTITY_HEADER, outline=ENTITY_HEADER)
    draw.text((x + 18, y + 9), title, fill=INK, font=FONT_HEADER)
    current_y = y + 64
    for line_text in lines:
        draw.text((x + 18, current_y), line_text, fill=INK, font=FONT_BODY)
        current_y += 28


def draw_flow_arrow(draw: ImageDraw.ImageDraw, x: int, y1: int, y2: int, label: str) -> None:
    """Draw a vertical arrow in the flowchart."""

    draw.line((x, y1, x, y2), fill=LINE, width=5)
    draw.polygon([(x, y2), (x - 10, y2 - 18), (x + 10, y2 - 18)], fill=LINE)
    bbox = draw.textbbox((0, 0), label, font=FONT_LABEL)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    label_x = x + 22
    label_y = (y1 + y2) // 2 - text_h // 2
    draw.rounded_rectangle(
        (label_x - 8, label_y - 8, label_x + text_w + 8, label_y + text_h + 8),
        radius=10,
        fill=LABEL_FILL,
        outline=ACCENT,
        width=2,
    )
    draw.text((label_x, label_y - 1), label, fill=ACCENT, font=FONT_LABEL)


def draw_flow_side_box(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, title: str, lines: list[str]) -> None:
    """Draw one side note box."""

    rounded_box(draw, x, y, w, h, fill=ENTITY_FILL, outline=ENTITY_BORDER)
    draw.rounded_rectangle((x, y, x + w, y + 42), radius=18, fill=ENTITY_HEADER, outline=ENTITY_BORDER, width=3)
    draw.rectangle((x, y + 20, x + w, y + 42), fill=ENTITY_HEADER, outline=ENTITY_HEADER)
    draw.text((x + 14, y + 7), title, fill=INK, font=FONT_LABEL)
    current_y = y + 58
    for line_text in lines:
        draw.text((x + 14, current_y), line_text, fill=INK, font=FONT_SMALL)
        current_y += 26


def generate_flowchart() -> None:
    """Generate the business flowchart image."""

    width, height = 2400, 2650
    image = Image.new("RGB", (width, height), BG)
    draw = ImageDraw.Draw(image)

    draw.text((70, 35), "学习效果评估与个性化反馈智能体业务流程图", fill=INK, font=FONT_TITLE)
    draw.text((70, 92), "Business Flowchart for Evaluation and Personalized Feedback Agent", fill=MUTED, font=FONT_SUB)

    center_x = 1200
    box_w = 1260
    box_h = 150
    start_x = center_x - box_w // 2
    positions = [180, 390, 600, 840, 1060, 1280, 1500, 1720, 1940, 2160]

    boxes = [
        ("1. 数据来源与接口入口", [
            "来源：前端答题提交、user-service、兼容接口 /evaluation/practice/submit、/evaluation/mistakes/qa",
            "入口：POST /evaluation/submit、/evaluation/batch_submit，以及阶段报告 / 月报查询接口",
        ]),
        ("2. 参数校验与请求归一化", [
            "校验 user_id、exercise_id、exercise_type、difficulty、time_spent、knowledge_point_ids",
            "统一题型别名、难度枚举、章节信息，并解析主观题 reference_answer / max_score",
        ]),
        ("3. 题目元数据读取与补建", [
            "查询 exercises、knowledge_points，读取题目内容、标准答案、章节信息",
            "若题目或知识点不存在，则自动 upsert，保证后续评估基于真实数据落库",
        ]),
        ("4. 分流判分：客观题 / 主观题 / 编程题", [
            "客观题：直接比对 user_answer 与 standard_answer，得到 is_correct 与 ratio",
            "主观题 / 编程题：调用 LLM 按参考答案评分，得到 score、comment、suggestion",
        ]),
        ("5. 汇总即时反馈结果", [
            "统一生成 is_correct、score、ratio、comment、suggestion、encouragement、error_pattern",
            "形成可直接返回前端的即时反馈结果，并为后续画像更新准备结构化数据",
        ]),
        ("6. 持久化答题事实数据", [
            "写入 answer_records：user_answer、is_correct、time_spent、created_at",
            "在 evaluation_json 中保存评分明细、章节信息、错误模式、建议、是否使用 LLM",
        ]),
        ("7. 更新用户画像与掌握度", [
            "基于知识点维度更新 user_profiles.mastery_json：掌握度、历史轨迹、正确率、连续错误次数",
            "同步更新 habits、cognitive_abilities、最近反馈摘要与最近作答时间",
        ]),
        ("8. 薄弱点识别与错误模式分析", [
            "识别规则：连续错 3 次，或近 10 题正确率低于阈值；同时聚合概念混淆、计算错误、推理不完整等模式",
            "形成 weak_knowledge_points、dominant_error_pattern，并作为画像与路径调整依据",
        ]),
        ("9. 事件发送与报告生成", [
            "向 RabbitMQ 发送 ProfileUpdateEvent 与 PathAdjustmentRequest，并支持阶段报告、月报、画像快照、学习建议查询",
            "阶段报告 / 月报聚合 answer_records + user_profiles 后写入 learning_reports",
        ]),
        ("10. 结果回传与下游联动", [
            "前端获得即时反馈、阶段报告、月报、画像快照、学习建议；教师端可读取详细报告与错题本",
            "画像智能体、路径规划智能体、后续推荐模块继续消费评估结果形成闭环",
        ]),
    ]

    for index, (title, lines) in enumerate(boxes):
        draw_flow_box(draw, start_x, positions[index], box_w, box_h, title, lines)

    for index in range(len(positions) - 1):
        draw_flow_arrow(draw, center_x, positions[index] + box_h, positions[index + 1], "数据流")

    # Detail branch: objective scoring
    draw_flow_side_box(
        draw,
        120,
        805,
        360,
        180,
        "客观题判分分支",
        ["适用：choice / fill / judge", "直接比对标准答案", "输出：is_correct = True / False"],
    )
    draw.line((480, 895, start_x, 895), fill=LINE, width=4)
    draw.polygon([(start_x, 895), (start_x + 16, 887), (start_x + 16, 903)], fill=LINE)
    draw_cardinality(draw, (550, 862), "客观题")

    # Detail branch: subjective scoring
    draw_flow_side_box(
        draw,
        1920,
        805,
        360,
        180,
        "主观题 / 编程题分支",
        ["适用：short_answer / code", "调用 LLM 严格评分", "输出：score、comment、suggestion"],
    )
    draw.line((start_x + box_w, 895, 1920, 895), fill=LINE, width=4)
    draw.polygon([(1920, 895), (1904, 887), (1904, 903)], fill=LINE)
    draw_cardinality(draw, (1845, 862), "主观题 / 编程题")

    # Data store boxes
    draw_flow_side_box(
        draw,
        90,
        1115,
        320,
        190,
        "题目主数据",
        ["exercises", "标准答案 / 参考答案", "题型、难度、章节信息"],
    )
    draw.line((410, 1210, start_x, 1210), fill=LINE, width=4)
    draw.polygon([(start_x, 1210), (start_x + 16, 1202), (start_x + 16, 1218)], fill=LINE)
    draw_cardinality(draw, (500, 1178), "读取 / 补建")

    draw_flow_side_box(
        draw,
        90,
        1450,
        320,
        190,
        "事实表",
        ["answer_records", "真实答题行为数据", "后续报告与分析基础"],
    )
    draw.line((410, 1545, start_x, 1545), fill=LINE, width=4)
    draw.polygon([(start_x, 1545), (start_x + 16, 1537), (start_x + 16, 1553)], fill=LINE)
    draw_cardinality(draw, (500, 1512), "落库")

    draw_flow_side_box(
        draw,
        1950,
        1670,
        330,
        200,
        "状态表",
        ["user_profiles", "掌握度、弱点标签", "画像习惯与认知摘要"],
    )
    draw.line((start_x + box_w, 1760, 1950, 1760), fill=LINE, width=4)
    draw.polygon([(1950, 1760), (1934, 1752), (1934, 1768)], fill=LINE)
    draw_cardinality(draw, (1875, 1726), "状态更新")

    draw_flow_side_box(
        draw,
        1950,
        1920,
        330,
        200,
        "消息队列",
        ["RabbitMQ", "ProfileUpdateEvent", "PathAdjustmentRequest"],
    )
    draw.line((start_x + box_w, 2015, 1950, 2015), fill=LINE, width=4)
    draw.polygon([(1950, 2015), (1934, 2007), (1934, 2023)], fill=LINE)
    draw_cardinality(draw, (1862, 1982), "事件发送")

    draw_flow_side_box(
        draw,
        1950,
        2185,
        330,
        220,
        "报告与衍生输出",
        ["learning_reports", "阶段报告 / 月报", "画像快照 / 学习建议"],
    )
    draw.line((start_x + box_w, 2260, 1950, 2260), fill=LINE, width=4)
    draw.polygon([(1950, 2260), (1934, 2252), (1934, 2268)], fill=LINE)
    draw_cardinality(draw, (1848, 2228), "查询 / 写入")

    # External consumers
    draw_flow_side_box(
        draw,
        90,
        2210,
        360,
        210,
        "下游消费者",
        ["前端学生端", "教师端洞察页", "画像智能体 / 路径规划智能体"],
    )
    draw.line((450, 2305, start_x, 2305), fill=LINE, width=4)
    draw.polygon([(start_x, 2305), (start_x + 16, 2297), (start_x + 16, 2313)], fill=LINE)
    draw_cardinality(draw, (560, 2272), "结果消费")

    rounded_box(draw, 90, 2460, 2220, 130, fill=LEGEND_FILL, outline=ENTITY_BORDER)
    draw.text(
        (110, 2478),
        "说明：该流程图在主链路之外补充展示了客观题 / 主观题分支、数据库读写点、RabbitMQ 事件发送，以及阶段报告、月报、画像快照、学习建议等衍生输出。",
        fill=MUTED,
        font=FONT_SMALL,
    )
    draw.text(
        (110, 2516),
        "其中 exercises 负责题目元数据，answer_records 是事实数据中心，user_profiles 是分析状态中心，learning_reports 负责结果沉淀，下游智能体继续消费评估结果形成闭环。",
        fill=MUTED,
        font=FONT_SMALL,
    )

    image.save(FLOW_IMAGE)


def configure_styles(document: Document) -> None:
    """Configure Word styles."""

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if style_name == "Normal":
            style.font.size = Pt(10.5)
    document.styles["Heading 2"].font.size = Pt(14)
    document.styles["Heading 2"].font.bold = True


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add one table to the document."""

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_document() -> None:
    """Build the final Word document."""

    document = Document()
    configure_styles(document)

    p = document.add_paragraph(style="Normal")
    p.alignment = 1
    run = p.add_run("AI 时代软件开发课程设计")
    run.bold = True
    run.font.size = Pt(30)

    p = document.add_paragraph(style="Normal")
    p.alignment = 1
    run = p.add_run("个人设计文档")
    run.font.size = Pt(18)

    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph(style="Normal")
    run = p.add_run("文档说明")
    run.bold = True
    run.font.size = Pt(13)

    document.add_paragraph(
        "本文档参照课程设计需求分析文档的排版方式，整理当前系统中的数据库表结构、实体关系图以及学习效果评估与个性化反馈智能体业务流程图。",
        style="Normal",
    )
    document.add_paragraph(
        "文档内容围绕“数据库表设计 → E-R 关系说明 → 智能体业务流程”展开，便于课程设计展示、论文撰写与后续系统实现。",
        style="Normal",
    )

    document.add_paragraph("")
    document.add_paragraph("1.1  核心数据库表", style="Heading 2")
    document.add_paragraph(
        "系统数据库的核心表主要包括用户、画像、知识点、题目、答题记录、学习报告、学习路径及任务等。各表共同支撑了学习资源生成、学习路径规划、答题评估与个性化反馈等智能体功能。",
        style="Normal",
    )
    add_table(
        document,
        ["表名", "作用", "主键", "关键字段"],
        [
            ["users", "平台用户基础信息", "id", "username, role, email, created_at"],
            ["user_profiles", "用户画像与知识掌握度", "user_id", "mastery_json, cognitive_abilities, habits, updated_at"],
            ["knowledge_points", "知识点定义表", "id", "name, description, difficulty, importance"],
            ["knowledge_relations", "知识点依赖关系", "id", "from_id, to_id, relation_type"],
            ["resources", "学习资源表", "id", "type, content, knowledge_point_id, generated_for_user_id"],
            ["learning_paths", "个性化学习路径", "id", "user_id, path_data_json, status, created_at"],
            ["learning_tasks", "学习路径任务明细", "id", "path_id, task_type, resource_ids, completed"],
            ["exercises", "题目主表", "id", "knowledge_point_id, type, difficulty, content, answer, analysis"],
            ["answer_records", "用户答题记录主表", "id", "user_id, exercise_id, user_answer, is_correct, time_spent, evaluation_json, created_at"],
            ["learning_reports", "学习报告表", "id", "user_id, report_type, content_json, created_at"],
            ["profile_conversations", "画像对话历史", "id", "user_id, role, content, created_at"],
        ],
    )

    document.add_paragraph("")
    document.add_paragraph("1.2  评估智能体关键表字段", style="Heading 2")
    document.add_paragraph(
        "在学习效果评估与个性化反馈智能体中，最关键的表为 answer_records、user_profiles 与 learning_reports。它们分别承担真实答题数据存储、掌握度状态维护以及分析结果沉淀的职责。",
        style="Normal",
    )
    key_sections = [
        (
            "answer_records",
            [
                ["id", "答题记录唯一标识"],
                ["user_id", "作答用户 ID"],
                ["exercise_id", "对应题目 ID"],
                ["user_answer", "用户原始答案"],
                ["is_correct", "最终判定是否正确"],
                ["time_spent", "作答耗时，单位秒"],
                ["evaluation_json", "评分明细、建议、章节信息、错误模式、是否使用 LLM"],
                ["created_at", "答题提交时间"],
            ],
        ),
        (
            "user_profiles",
            [
                ["user_id", "画像主键，同时关联 users"],
                ["mastery_json", "知识点掌握度、最近正确率、连续错误、成长历史、弱点标签"],
                ["learning_style", "学习风格标签"],
                ["cognitive_abilities", "错误模式摘要、薄弱知识点集合、最近反馈"],
                ["habits", "累计答题数、平均耗时、最近作答时间等"],
                ["updated_at", "画像更新时间"],
            ],
        ),
        (
            "learning_reports",
            [
                ["id", "报告唯一标识"],
                ["user_id", "报告所属用户"],
                ["report_type", "阶段报告、月报等类型"],
                ["content_json", "报告完整内容"],
                ["created_at", "报告生成时间"],
            ],
        ),
    ]

    for title, rows in key_sections:
        document.add_paragraph(f"{title}", style="Normal")
        add_table(document, ["字段名", "说明"], rows)
        document.add_paragraph("")

    document.add_paragraph("1.3  E-R 关系表", style="Heading 2")
    document.add_paragraph(
        "数据库各实体之间的关系主要围绕“用户—题目—答题记录—画像—报告—知识点—学习路径”展开。下表对核心实体间的关系进行说明。",
        style="Normal",
    )
    add_table(
        document,
        ["主实体", "关系", "从实体", "说明"],
        [
            ["users", "1:1", "user_profiles", "一个用户对应一份画像"],
            ["users", "1:N", "answer_records", "一个用户有多条答题记录"],
            ["exercises", "1:N", "answer_records", "一道题可被多次作答"],
            ["knowledge_points", "1:N", "exercises", "一个知识点可关联多道题"],
            ["users", "1:N", "learning_reports", "一个用户可生成多份报告"],
            ["users", "1:N", "learning_paths", "一个用户可拥有多条学习路径"],
            ["learning_paths", "1:N", "learning_tasks", "一条路径包含多个任务"],
            ["knowledge_points", "N:N", "knowledge_relations", "知识点之间形成依赖关系图"],
            ["knowledge_points", "1:N", "resources", "一个知识点可关联多个资源"],
        ],
    )

    document.add_paragraph("")
    document.add_paragraph("1.4  系统 E-R 图", style="Heading 2")
    document.add_paragraph(
        "下图展示了学习效果评估与个性化反馈智能体涉及的核心实体、主外键关系以及 1:1、1:N、N:N 的主要结构。",
        style="Normal",
    )
    document.add_picture(str(ER_IMAGE), width=Inches(7.8))
    document.add_paragraph("图 1  学习效果评估与个性化反馈智能体 E-R 图")

    document.add_page_break()
    document.add_paragraph("1.5  智能体业务流程图", style="Heading 2")
    document.add_paragraph(
        "下图展示了学习效果评估与个性化反馈智能体从答题提交、题目解析、判分反馈，到答题落库、画像更新、事件发送和报告生成的完整业务闭环。",
        style="Normal",
    )
    document.add_picture(str(FLOW_IMAGE), width=Inches(7.8))
    document.add_paragraph("图 2  学习效果评估与个性化反馈智能体业务流程图")

    document.save(OUTPUT_DOCX)


def main() -> None:
    """Generate all assets and verify embedded images."""

    generate_er_diagram()
    generate_flowchart()
    build_document()

    with ZipFile(OUTPUT_DOCX, "r") as zip_file:
        media_files = [name for name in zip_file.namelist() if name.startswith("word/media/")]

    print(OUTPUT_DOCX)
    print(f"embedded_media_count={len(media_files)}")
    for name in media_files:
        print(name)


if __name__ == "__main__":
    main()
